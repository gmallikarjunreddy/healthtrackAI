"""
HealthTrackAI — File Parser
Extracts text from lab reports using:
  PDF  → pdfplumber (text-based) + Tesseract OCR (scanned/image PDFs)
  IMG  → pytesseract (Tesseract OCR) — 100% free, no API needed
  DOCX → python-docx
  TXT  → direct decode
"""
import io
import os
import re
from pathlib import Path

import numpy as np
import requests
from PIL import Image


# ── Tesseract setup ────────────────────────────────────────────────────────────
try:
    import pytesseract
    # Windows path override from .env
    tess_cmd = os.getenv("TESSERACT_CMD", "")
    if tess_cmd:
        pytesseract.pytesseract.tesseract_cmd = tess_cmd
    TESSERACT_OK = True
except ImportError:
    TESSERACT_OK = False


def _ocr_image(pil_img: Image.Image) -> str:
    """Run Tesseract OCR on a PIL image with preprocessing for better accuracy."""
    if not TESSERACT_OK:
        return "[Tesseract not installed. Run: pip install pytesseract and install Tesseract binary]"
    try:
        import cv2
        # Convert to grayscale numpy array
        img_arr = np.array(pil_img.convert("RGB"))
        gray = cv2.cvtColor(img_arr, cv2.COLOR_RGB2GRAY)
        # Denoise + threshold for cleaner OCR
        denoised = cv2.fastNlMeansDenoising(gray, h=10)
        _, thresh = cv2.threshold(denoised, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
        # Scale up small images for better OCR
        h, w = thresh.shape
        if h < 600:
            scale = 600 / h
            thresh = cv2.resize(thresh, (int(w * scale), int(h * scale)), interpolation=cv2.INTER_CUBIC)
        result_img = Image.fromarray(thresh)
        text = pytesseract.image_to_string(
            result_img,
            config="--psm 6 --oem 3 -c tessedit_char_whitelist=0123456789ABCDEFGHIJKLMNOPQRSTUVWXYZabcdefghijklmnopqrstuvwxyz.,:-/()% "
        )
        return text.strip()
    except Exception as e:
        # Fallback: plain Tesseract without preprocessing
        try:
            return pytesseract.image_to_string(pil_img).strip()
        except Exception as e2:
            return f"[OCR failed: {e2}]"


def extract_text_from_file(file_bytes: bytes, filename: str) -> str:
    """Extract text from any supported file type."""
    ext = Path(filename).suffix.lower()
    try:
        if ext == ".pdf":
            return _extract_pdf(file_bytes)
        elif ext in (".png", ".jpg", ".jpeg", ".bmp", ".tiff", ".webp"):
            return _extract_image(file_bytes)
        elif ext == ".docx":
            return _extract_docx(file_bytes)
        elif ext == ".txt":
            return file_bytes.decode("utf-8", errors="ignore")
        else:
            return f"[Unsupported format '{ext}'. Use PDF, JPG, PNG, DOCX, or TXT]"
    except Exception as e:
        return f"[Parse error for '{filename}': {e}]"


def _extract_pdf(data: bytes) -> str:
    import pdfplumber
    text_parts = []
    ocr_pages = []

    with pdfplumber.open(io.BytesIO(data)) as pdf:
        for i, page in enumerate(pdf.pages):
            # Try text extraction first
            t = page.extract_text()
            if t and len(t.strip()) > 30:
                text_parts.append(t)
                # Extract tables
                for table in page.extract_tables():
                    for row in table:
                        if row:
                            cleaned = " | ".join(str(c).strip() for c in row if c and str(c).strip())
                            if cleaned:
                                text_parts.append(cleaned)
            else:
                # Scanned PDF page → use Tesseract OCR
                ocr_pages.append(i)
                if TESSERACT_OK:
                    try:
                        pil_img = page.to_image(resolution=200).original
                        ocr_text = _ocr_image(pil_img)
                        if ocr_text:
                            text_parts.append(f"[Page {i+1} — OCR]\n{ocr_text}")
                    except Exception:
                        pass

    if not text_parts:
        return "[PDF appears empty. Ensure the file contains readable text or a clear scan.]"

    result = "\n".join(text_parts)
    if ocr_pages:
        result = f"[Note: Pages {ocr_pages} were scanned — Tesseract OCR applied]\n\n" + result
    return result


def _extract_image(data: bytes) -> str:
    """Extract text from image-based lab report via Tesseract."""
    pil_img = Image.open(io.BytesIO(data))
    text = _ocr_image(pil_img)
    if not text or len(text) < 20:
        return "[Image OCR returned minimal text. Ensure the image is clear, high-resolution, and well-lit.]"
    return text


def _extract_docx(data: bytes) -> str:
    from docx import Document
    doc = Document(io.BytesIO(data))
    parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def extract_from_url(url: str) -> tuple[bytes, str, str]:
    """Download an image from a URL and return (bytes, filename, error)."""
    try:
        if not url.startswith(("http://", "https://")):
            return b"", "", "Invalid URL format. Use http:// or https://"
        
        resp = requests.get(url, timeout=12, stream=True)
        resp.raise_for_status()

        # Try to get filename from content-disposition
        cd = resp.headers.get("Content-Disposition", "")
        if "filename=" in cd:
            fname = re.findall("filename=(.+)", cd)[0].strip('"')
        else:
            # Try to get filename from URL
            fname = url.split("/")[-1].split("?")[0]
            if not fname or "." not in fname:
                # Fallback based on content-type
                ctype = resp.headers.get("content-type", "").lower()
                ext = ".pdf" if "pdf" in ctype else ".jpg"
                if "png" in ctype: ext = ".png"
                elif "bmp" in ctype: ext = ".bmp"
                elif "webp" in ctype: ext = ".webp"
                elif "docx" in ctype: ext = ".docx"
                elif "text" in ctype or "plain" in ctype: ext = ".txt"
                
                from time import time
                fname = f"downloaded_report_{int(time())}{ext}"
        
        return resp.content, fname, ""
    except Exception as e:
        return b"", "", f"Download failed: {str(e)}"


def detect_report_type(filename: str, text: str) -> str:
    """Classify the type of lab report from content."""
    combined = (filename + " " + text[:800]).lower()
    checks = [
        (["hemoglobin", "hematocrit", "rbc", "wbc", "platelet", "cbc", "blood count"], "Blood Report (CBC)"),
        (["vitamin d", "vitamin b12", "b12", "folate", "ferritin", "vit d", "vitamin panel"], "Vitamin Panel"),
        (["tsh", "t3", "t4", "thyroid", "thyroxine"], "Thyroid Function"),
        (["ldl", "hdl", "triglyceride", "cholesterol", "lipid"], "Lipid Panel"),
        (["mri", "ct scan", "x-ray", "xray", "radiology", "imaging", "ultrasound"], "Imaging Report"),
        (["urine", "urinalysis", "creatinine", "urea", "uric acid", "protein in urine"], "Urine Analysis"),
        (["hba1c", "glucose", "fasting sugar", "diabetes", "insulin"], "Diabetes Panel"),
        (["sgpt", "sgot", "alt", "ast", "bilirubin", "liver function", "lft"], "Liver Function"),
        (["sodium", "potassium", "calcium", "electrolyte", "chloride"], "Electrolyte Panel"),
        (["kidney", "gfr", "egfr", "creatinine serum", "renal"], "Kidney Function"),
        (["cortisol", "hormone", "estrogen", "testosterone", "progesterone", "lh", "fsh"], "Hormone Panel"),
    ]
    for keywords, label in checks:
        if any(k in combined for k in keywords):
            return label
    return "General Lab Report"


def get_report_icon(report_type: str) -> str:
    icons = {
        "Blood Report (CBC)": "🩸",
        "Vitamin Panel": "💊",
        "Thyroid Function": "🦋",
        "Lipid Panel": "🫀",
        "Imaging Report": "🩻",
        "Urine Analysis": "🧪",
        "Diabetes Panel": "🍬",
        "Liver Function": "🫁",
        "Electrolyte Panel": "⚡",
        "Kidney Function": "🫘",
        "Hormone Panel": "⚗️",
        "General Lab Report": "📋",
    }
    return icons.get(report_type, "📋")
